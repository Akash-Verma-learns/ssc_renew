"""
core/document_parser.py
========================
Multi-Strategy Document Parser
================================

Parser cascade (auto-selects best available, tries in order):
  1. LlamaParse   — cloud API, best quality for complex PDFs with tables
  2. Mistral OCR  — vision-based OCR, excellent for scanned/image PDFs
  3. Docling      — local table-aware structured parser (TableFormer model)
  4. PyMuPDF      — always-available geometry-based fallback

All parsers return the same Chunk dataclass, fully compatible with
the existing vector_store.py ChromaDB pipeline.

Tables are preserved with:
  - Original markdown (| col1 | col2 | col3 |) structure
  - Cell relationships intact (no lost columns)
  - is_table flag for downstream filtering
  - table_markdown field for direct rendering

Environment variables (all optional):
  LLAMA_CLOUD_API_KEY — enables LlamaParse (best table extraction)
  MISTRAL_API_KEY     — enables Mistral OCR (best for scanned PDFs)
  PARSER_STRATEGY     — force: "llamaparse"|"mistral"|"docling"|"pymupdf"|"auto"
"""

from __future__ import annotations

import base64
import hashlib
import json
import os
import re
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Tuple

# ─────────────────────────────────────────────────────────────────────────────
# Chunk dataclass — backward-compatible with existing core/parser.py
# ─────────────────────────────────────────────────────────────────────────────

@dataclass
class Chunk:
    """Semantic text chunk from a parsed document."""
    text:            str
    page_no:         int
    section_heading: str   = ""
    clause_ref:      str   = ""
    doc_name:        str   = ""
    chunk_id:        str   = field(default="")
    # Enhanced fields (new — ignored by old code that doesn't check them)
    is_table:        bool  = False
    table_markdown:  str   = ""    # original markdown table for LLM prompts
    parser_used:     str   = "pymupdf"
    confidence:      float = 1.0   # 0.5 = OCR estimate, 1.0 = direct text

    def __post_init__(self):
        if not self.chunk_id:
            h = hashlib.md5(
                (self.doc_name + str(self.page_no) + self.text[:80]).encode()
            ).hexdigest()[:8]
            self.chunk_id = f"{self.doc_name}_{self.page_no}_{h}"


# ─────────────────────────────────────────────────────────────────────────────
# Config
# ─────────────────────────────────────────────────────────────────────────────

LLAMA_CLOUD_API_KEY = os.getenv("LLAMA_CLOUD_API_KEY", "")
MISTRAL_API_KEY     = os.getenv("MISTRAL_API_KEY", "")
PARSER_STRATEGY     = os.getenv("PARSER_STRATEGY", "auto")

MIN_CHUNK_CHARS = 40
MAX_CHUNK_CHARS = 1800

# Heading patterns for section detection
_HEADING_PATTERNS = [
    re.compile(r"^(\d+)\.\s+[A-Z]"),          # "1. Introduction"
    re.compile(r"^(\d+\.\d+)\s+[A-Z]"),        # "1.1 Background"
    re.compile(r"^(\d+\.\d+\.\d+)\s+[A-Z]"),   # "1.1.1 Context"
    re.compile(r"^(Clause|Section|Article|Part|Schedule)\s+\d+", re.I),
    re.compile(r"^[A-Z][A-Z\s]{5,60}$"),        # ALL CAPS HEADINGS
    re.compile(r"^(SCOPE|PAYMENT|LIABILITY|INSURANCE|TERMINATION)", re.I),
    re.compile(r"^Form\s+(Tech|Fin)-\d+", re.I),  # "Form Tech-1"
    re.compile(r"^S\.\s*No\.", re.I),             # Table header
]


def _is_heading(text: str, font_size: float = None, is_bold: bool = None) -> bool:
    text = text.strip()
    if not text or len(text) > 250:
        return False
    if font_size and font_size > 12.5:
        return True
    if is_bold and len(text) < 150:
        return True
    return any(p.match(text) for p in _HEADING_PATTERNS)


def _extract_clause_ref(heading: str) -> str:
    if not heading:
        return ""
    for p in [
        r"(clause\s+[\d.]+)",
        r"(section\s+[\d.]+)",
        r"(article\s+[\d.]+)",
        r"([\d]+\.[\d]+(?:\.[\d]+)*)",
    ]:
        m = re.search(p, heading, re.I)
        if m:
            return m.group(1).strip()
    return heading[:60]


# ─────────────────────────────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────────────────────────────

def parse_document(
    file_path: str,
    doc_name:  str = "",
    verbose:   bool = True,
) -> Tuple[List[Chunk], str]:
    """
    Parse a document using the best available strategy.

    Returns:
        (chunks, parser_name_used)

    The cascade tries parsers in priority order and uses the first that succeeds.
    Falls back gracefully if API keys are missing or parsers fail.
    """
    file_path = str(file_path)
    doc_name  = doc_name or Path(file_path).name
    ext       = Path(file_path).suffix.lower()

    if verbose:
        print(f"[Parser] Parsing '{Path(file_path).name}' (strategy={PARSER_STRATEGY})")

    # DOCX: python-docx is always the best choice
    if ext in (".docx", ".doc"):
        chunks = _parse_docx(file_path, doc_name)
        if verbose:
            print(f"[Parser] python-docx: {len(chunks)} chunks")
        return chunks, "python-docx"

    # PDF cascade
    strategy = PARSER_STRATEGY.lower()

    # 1. LlamaParse (best quality for complex PDFs)
    if LLAMA_CLOUD_API_KEY and strategy in ("auto", "llamaparse"):
        try:
            chunks = _parse_llamaparse(file_path, doc_name)
            if chunks:
                if verbose:
                    tables = sum(1 for c in chunks if c.is_table)
                    print(f"[Parser] LlamaParse: {len(chunks)} chunks, {tables} tables ✓")
                return chunks, "llamaparse"
        except Exception as e:
            if verbose:
                print(f"[Parser] LlamaParse failed: {e} — trying next")

    # 2. Mistral OCR (excellent for scanned/complex PDFs)
    if MISTRAL_API_KEY and strategy in ("auto", "mistral"):
        try:
            chunks = _parse_mistral_ocr(file_path, doc_name)
            if chunks:
                if verbose:
                    tables = sum(1 for c in chunks if c.is_table)
                    print(f"[Parser] Mistral OCR: {len(chunks)} chunks, {tables} tables ✓")
                return chunks, "mistral-ocr"
        except Exception as e:
            if verbose:
                print(f"[Parser] Mistral OCR failed: {e} — trying next")

    # 3. Docling (local, table-aware)
    if strategy in ("auto", "docling"):
        try:
            chunks = _parse_docling(file_path, doc_name)
            if chunks:
                if verbose:
                    tables = sum(1 for c in chunks if c.is_table)
                    print(f"[Parser] Docling: {len(chunks)} chunks, {tables} tables ✓")
                return chunks, "docling"
        except ImportError:
            if verbose:
                print("[Parser] Docling not installed — trying PyMuPDF")
        except Exception as e:
            if verbose:
                print(f"[Parser] Docling failed: {e} — trying PyMuPDF")

    # 4. PyMuPDF (always available)
    chunks = _parse_pymupdf(file_path, doc_name)
    if verbose:
        tables = sum(1 for c in chunks if c.is_table)
        print(f"[Parser] PyMuPDF: {len(chunks)} chunks, {tables} table-like blocks ✓")
    return chunks, "pymupdf"


def parse_document_simple(file_path: str, doc_name: str = "") -> List[Chunk]:
    """Compatibility shim — returns just chunks (no parser name)."""
    chunks, _ = parse_document(file_path, doc_name, verbose=False)
    return chunks


# ─────────────────────────────────────────────────────────────────────────────
# Strategy 1: LlamaParse
# ─────────────────────────────────────────────────────────────────────────────

def _parse_llamaparse(file_path: str, doc_name: str) -> List[Chunk]:
    """
    LlamaParse: best-in-class PDF parser with accurate table extraction.
    Converts tables to markdown with cell relationships preserved.

    Requires: pip install llama-parse
    Env:      LLAMA_CLOUD_API_KEY
    """
    from llama_parse import LlamaParse
    from llama_index.core import SimpleDirectoryReader

    parser = LlamaParse(
        api_key=LLAMA_CLOUD_API_KEY,
        result_type="markdown",   # preserves table structure
        verbose=False,
        language="en",
        skip_diagonal_text=True,
        invalidate_cache=False,
    )

    file_extractor = {".pdf": parser}
    documents = SimpleDirectoryReader(
        input_files=[file_path],
        file_extractor=file_extractor,
    ).load_data()

    chunks: List[Chunk] = []
    current_heading  = ""
    chunk_idx        = 0

    for doc_obj in documents:
        # LlamaParse returns markdown per page
        page_no  = int(doc_obj.metadata.get("page_label", "0") or "0")
        markdown = doc_obj.text

        # Split on headings and tables
        sections = _split_markdown_sections(markdown)
        for section_text, is_table in sections:
            section_text = section_text.strip()
            if len(section_text) < MIN_CHUNK_CHARS:
                continue

            # Detect heading
            first_line = section_text.split("\n")[0].lstrip("#").strip()
            if first_line and (section_text.startswith("#") or _is_heading(first_line)):
                current_heading = first_line

            chunk_idx += 1
            table_md = section_text if is_table else ""
            # Convert table markdown to plain text for embedding
            plain = _table_to_plain(section_text) if is_table else section_text

            chunks.append(Chunk(
                text            = plain[:MAX_CHUNK_CHARS],
                page_no         = page_no,
                section_heading = current_heading,
                clause_ref      = _extract_clause_ref(current_heading),
                doc_name        = doc_name,
                chunk_id        = f"{doc_name}_{page_no}_{chunk_idx:04d}",
                is_table        = is_table,
                table_markdown  = table_md[:3000],
                parser_used     = "llamaparse",
                confidence      = 1.0,
            ))

    return chunks


def _split_markdown_sections(markdown: str) -> List[Tuple[str, bool]]:
    """
    Split LlamaParse markdown output into (text, is_table) pairs.
    Tables are identified by the | column | separator pattern.
    """
    sections: List[Tuple[str, bool]] = []
    lines    = markdown.split("\n")
    buf      = []
    in_table = False

    for line in lines:
        is_table_line = line.strip().startswith("|")
        if is_table_line != in_table:
            if buf:
                sections.append(("\n".join(buf), in_table))
                buf = []
            in_table = is_table_line
        buf.append(line)

    if buf:
        sections.append(("\n".join(buf), in_table))

    return sections


def _table_to_plain(table_md: str) -> str:
    """Convert markdown table to plain text for embedding."""
    lines = []
    for line in table_md.split("\n"):
        if "|" in line and not re.match(r"^\s*\|[-:]+\|", line):
            cells = [c.strip() for c in line.split("|") if c.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


# ─────────────────────────────────────────────────────────────────────────────
# Strategy 2: Mistral OCR
# ─────────────────────────────────────────────────────────────────────────────

def _parse_mistral_ocr(file_path: str, doc_name: str) -> List[Chunk]:
    """
    Mistral OCR: vision-based document understanding.
    Excellent for scanned PDFs, mixed image/text, and complex layouts.

    Requires: pip install mistralai
    Env:      MISTRAL_API_KEY
    """
    from mistralai import Mistral

    client = Mistral(api_key=MISTRAL_API_KEY)

    # Read and encode PDF
    with open(file_path, "rb") as f:
        pdf_bytes = f.read()
    b64_pdf = base64.b64encode(pdf_bytes).decode("utf-8")

    # Call Mistral OCR API
    response = client.ocr.process(
        model="mistral-ocr-latest",
        document={
            "type":          "document_url",
            "document_url":  f"data:application/pdf;base64,{b64_pdf}",
        },
        include_image_base64=False,
    )

    chunks: List[Chunk] = []
    current_heading = ""
    chunk_idx       = 0

    for page_result in (response.pages or []):
        page_no  = getattr(page_result, "index", 0) + 1
        markdown = getattr(page_result, "markdown", "") or ""

        sections = _split_markdown_sections(markdown)
        for section_text, is_table in sections:
            section_text = section_text.strip()
            if len(section_text) < MIN_CHUNK_CHARS:
                continue

            first_line = section_text.split("\n")[0].lstrip("#").strip()
            if first_line and _is_heading(first_line):
                current_heading = first_line

            chunk_idx += 1
            table_md = section_text if is_table else ""
            plain    = _table_to_plain(section_text) if is_table else section_text

            chunks.append(Chunk(
                text            = plain[:MAX_CHUNK_CHARS],
                page_no         = page_no,
                section_heading = current_heading,
                clause_ref      = _extract_clause_ref(current_heading),
                doc_name        = doc_name,
                chunk_id        = f"{doc_name}_{page_no}_{chunk_idx:04d}",
                is_table        = is_table,
                table_markdown  = table_md[:3000],
                parser_used     = "mistral-ocr",
                confidence      = 0.95,
            ))

    return chunks


# ─────────────────────────────────────────────────────────────────────────────
# Strategy 3: Docling (local, table-aware)
# ─────────────────────────────────────────────────────────────────────────────

def _parse_docling(file_path: str, doc_name: str) -> List[Chunk]:
    """
    Docling: uses TableFormer ML model for accurate table cell detection.
    Runs locally — no API key required. Best local option for tables.

    Requires: pip install docling
    """
    from docling.document_converter import DocumentConverter, PdfFormatOption
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions, TableFormerMode

    opts = PdfPipelineOptions()
    opts.do_table_structure             = True
    opts.do_ocr                         = False
    opts.table_structure_options.mode   = TableFormerMode.FAST
    opts.table_structure_options.do_cell_matching = True

    conv   = DocumentConverter(
        format_options={InputFormat.PDF: PdfFormatOption(pipeline_options=opts)}
    )
    result = conv.convert(file_path)
    doc    = result.document

    chunks: List[Chunk] = []
    current_heading     = ""
    chunk_idx           = 0

    # Text items
    for item, _ in doc.iterate_items():
        label   = str(getattr(item, "label", "")).lower()
        text    = getattr(item, "text",  "").strip()
        page_no = (item.prov[0].page_no if getattr(item, "prov", None) else 0)

        if not text or len(text) < MIN_CHUNK_CHARS:
            continue
        if "table" in label:
            continue  # handled separately below

        if "section_header" in label or "title" in label:
            current_heading = text

        chunk_idx += 1
        chunks.append(Chunk(
            text            = text[:MAX_CHUNK_CHARS],
            page_no         = page_no,
            section_heading = current_heading,
            clause_ref      = _extract_clause_ref(current_heading),
            doc_name        = doc_name,
            chunk_id        = f"{doc_name}_{page_no}_{chunk_idx:04d}",
            parser_used     = "docling",
            confidence      = 1.0,
        ))

    # Tables — converted to markdown with cell relationships preserved
    for tbl in doc.tables:
        page_no = (tbl.prov[0].page_no if getattr(tbl, "prov", None) else 0)
        try:
            md = tbl.export_to_markdown(doc)
        except Exception:
            md = ""
        if not md or len(md.strip()) < MIN_CHUNK_CHARS:
            continue

        plain = _table_to_plain(md)
        chunk_idx += 1
        chunks.append(Chunk(
            text            = f"[TABLE page {page_no}]\n{plain}"[:MAX_CHUNK_CHARS],
            page_no         = page_no,
            section_heading = current_heading,
            clause_ref      = f"p{page_no}_table",
            doc_name        = doc_name,
            chunk_id        = f"{doc_name}_{page_no}_t{chunk_idx:04d}",
            is_table        = True,
            table_markdown  = md[:3000],
            parser_used     = "docling",
            confidence      = 1.0,
        ))

    # Sort by page number
    chunks.sort(key=lambda c: c.page_no)
    return chunks


# ─────────────────────────────────────────────────────────────────────────────
# Strategy 4: PyMuPDF (always-available fallback)
# ─────────────────────────────────────────────────────────────────────────────

def _parse_pymupdf(file_path: str, doc_name: str) -> List[Chunk]:
    """
    PyMuPDF fallback: geometry-based text extraction.
    Fast and reliable for digital PDFs with embedded text.
    Tables are detected heuristically (column alignment scoring).
    """
    import fitz

    doc             = fitz.open(file_path)
    chunks          = []
    current_heading = "Preamble"
    chunk_idx       = 0

    for page_num, page in enumerate(doc, start=1):
        blocks = page.get_text("dict")["blocks"]
        page_buf: List[str] = []
        page_buf_bold       = False

        for block in blocks:
            if block["type"] != 0:
                continue
            for line in block["lines"]:
                line_text = " ".join(s["text"] for s in line["spans"]).strip()
                if not line_text:
                    continue

                max_size = max((s["size"] for s in line["spans"]), default=11)
                is_bold  = any(s["flags"] & (1 << 4) for s in line["spans"])

                if _is_heading(line_text, font_size=max_size, is_bold=is_bold):
                    # Flush buffer
                    if page_buf:
                        chunk_idx += 1
                        full = "\n".join(page_buf)
                        # Detect if this looks like a table (many | separators)
                        is_tbl = page_buf.count("|") > 3 if isinstance(page_buf, str) else \
                                 full.count("|") > 3
                        chunks.append(Chunk(
                            text            = full[:MAX_CHUNK_CHARS],
                            page_no         = page_num,
                            section_heading = current_heading,
                            clause_ref      = _extract_clause_ref(current_heading),
                            doc_name        = doc_name,
                            chunk_id        = f"{doc_name}_{page_num}_{chunk_idx:04d}",
                            is_table        = is_tbl,
                            parser_used     = "pymupdf",
                        ))
                        page_buf = []
                    current_heading = line_text
                    page_buf = [line_text]
                else:
                    page_buf.append(line_text)

                # Flush large buffers
                if sum(len(t) for t in page_buf) >= MAX_CHUNK_CHARS:
                    chunk_idx += 1
                    full = "\n".join(page_buf)
                    chunks.append(Chunk(
                        text            = full[:MAX_CHUNK_CHARS],
                        page_no         = page_num,
                        section_heading = current_heading,
                        clause_ref      = _extract_clause_ref(current_heading),
                        doc_name        = doc_name,
                        chunk_id        = f"{doc_name}_{page_num}_{chunk_idx:04d}",
                        parser_used     = "pymupdf",
                    ))
                    page_buf = []

        # Flush page buffer
        if page_buf and sum(len(t) for t in page_buf) >= MIN_CHUNK_CHARS:
            chunk_idx += 1
            full = "\n".join(page_buf)
            chunks.append(Chunk(
                text            = full[:MAX_CHUNK_CHARS],
                page_no         = page_num,
                section_heading = current_heading,
                clause_ref      = _extract_clause_ref(current_heading),
                doc_name        = doc_name,
                chunk_id        = f"{doc_name}_{page_num}_{chunk_idx:04d}",
                parser_used     = "pymupdf",
            ))

    doc.close()

    # Also extract table-like regions using column alignment
    _enrich_with_table_detection(chunks, file_path, doc_name)
    return chunks


def _enrich_with_table_detection(chunks: List[Chunk], file_path: str, doc_name: str):
    """
    Post-process PyMuPDF chunks to detect real tables using word bounding boxes.
    Only marks chunks as tables if they have clear multi-column structure.
    """
    try:
        import fitz
        doc = fitz.open(file_path)
        for chunk in chunks:
            # Reset false-positive table flags from initial detection
            chunk.is_table = False
            chunk.table_markdown = ""

            if chunk.page_no == 0 or chunk.page_no > len(doc):
                continue

            page  = doc[chunk.page_no - 1]
            words = page.get_text("words")
            if not words:
                continue

            # Only detect tables: need 3+ distinct X-column clusters AND row structure
            x_positions = [round(w[0] / 15) * 15 for w in words]
            from collections import Counter
            x_counts = Counter(x_positions)
            # At least 3 columns with multiple hits = likely table
            column_xs = [x for x, count in x_counts.items() if count >= 3]
            if len(column_xs) < 3:
                continue

            # Also need rows (multiple Y positions)
            y_positions = sorted(set(round(w[1] / 10) * 10 for w in words))
            if len(y_positions) < 4:
                continue

            # Check chunk text actually contains table-like content
            chunk_text_lower = chunk.text.lower()
            table_signals = ["s.no", "s. no", "marks", "criteria", "particulars",
                             "parameter", "qualification", "experience", "score",
                             "amount", "value", "rate", "units", "total"]
            has_signal = any(sig in chunk_text_lower for sig in table_signals)
            if not has_signal:
                continue

            # This chunk is on a page with real table structure
            md = _build_table_markdown(words, page.rect.width)
            if md and md.count("|") > 12:  # require substantial column structure
                chunk.is_table       = True
                chunk.table_markdown = md[:3000]

        doc.close()
    except Exception:
        pass


def _build_table_markdown(words: list, page_width: float) -> str:
    """Build a rough markdown table from word bounding boxes."""
    if not words:
        return ""
    # Group words by Y position (rows)
    rows: Dict[int, list] = {}
    for x0, y0, x1, y1, text, *_ in words:
        row_key = round(y0 / 8) * 8
        rows.setdefault(row_key, []).append((x0, text))

    lines = []
    for y in sorted(rows)[:30]:  # max 30 rows
        cells = sorted(rows[y], key=lambda w: w[0])
        line  = " | ".join(c[1] for c in cells)
        lines.append(f"| {line} |")

    return "\n".join(lines)


# ─────────────────────────────────────────────────────────────────────────────
# DOCX parser
# ─────────────────────────────────────────────────────────────────────────────

def _parse_docx(file_path: str, doc_name: str) -> List[Chunk]:
    """Parse DOCX using python-docx. Tables are extracted as markdown."""
    from docx import Document as DocxDocument

    docx_doc    = DocxDocument(file_path)
    chunks      = []
    para_count  = 0
    current_h   = "Preamble"
    buf: List[str] = []
    chunk_idx   = 0
    PARAS_PER_PAGE = 35

    def _flush(pg: int):
        nonlocal chunk_idx
        text = "\n".join(buf).strip()
        if len(text) >= MIN_CHUNK_CHARS:
            chunk_idx += 1
            chunks.append(Chunk(
                text            = text[:MAX_CHUNK_CHARS],
                page_no         = pg,
                section_heading = current_h,
                clause_ref      = _extract_clause_ref(current_h),
                doc_name        = doc_name,
                chunk_id        = f"{doc_name}_{pg}_{chunk_idx:04d}",
                parser_used     = "python-docx",
            ))
        buf.clear()

    for para in docx_doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        para_count += 1
        page_no = (para_count // PARAS_PER_PAGE) + 1
        is_h    = para.style.name.startswith("Heading") or _is_heading(text)
        if is_h:
            _flush(page_no)
            current_h = text
            buf.append(text)
        else:
            buf.append(text)

    _flush(para_count // PARAS_PER_PAGE + 1)

    # Extract tables as markdown
    for t_idx, table in enumerate(docx_doc.tables):
        rows_md = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            rows_md.append("| " + " | ".join(cells) + " |")
        if len(rows_md) >= 2:
            table_md = "\n".join(rows_md)
            plain    = _table_to_plain(table_md)
            chunk_idx += 1
            chunks.append(Chunk(
                text            = plain[:MAX_CHUNK_CHARS],
                page_no         = 0,
                section_heading = current_h,
                clause_ref      = f"table_{t_idx+1}",
                doc_name        = doc_name,
                chunk_id        = f"{doc_name}_table_{chunk_idx:04d}",
                is_table        = True,
                table_markdown  = table_md[:3000],
                parser_used     = "python-docx",
            ))

    return chunks
